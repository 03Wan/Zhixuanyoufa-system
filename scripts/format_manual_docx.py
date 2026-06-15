from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "manual_exact.md"
OUT_ASCII = ROOT / "Thesis_Formatted_Manual.docx"
OUT_CN = ROOT / "智选优发系统队员学习与比赛演示手册_论文风格版.docx"


def set_paragraph_format(p, first_line_indent: bool) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)


def clean_inline(s: str) -> str:
    s = s.replace("**", "").replace("`", "")
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", s)
    return s


def apply_font_run(run, east_asia: str = "宋体", size: float = 12.0, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def build_document(text: str) -> Document:
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name, east, size in [("Heading 1", "黑体", 16), ("Heading 2", "黑体", 14), ("Heading 3", "黑体", 12)]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        st.font.size = Pt(size)
        st.font.bold = True

    title = "智选优发系统队员学习与比赛演示手册"
    subtitle = "（正式论文风格排版版）"

    p = doc.add_paragraph(title)
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font_run(p.runs[0], east_asia="黑体", size=22, bold=True)

    p2 = doc.add_paragraph(subtitle)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font_run(p2.runs[0], east_asia="宋体", size=14)
    set_paragraph_format(p2, first_line_indent=False)

    for _ in range(6):
        doc.add_paragraph("")

    meta = [
        "项目名称：智选优发——基于AIGC的跨境商品素材智能评估与发布决策辅助系统",
        "文档用途：大学生创新创业比赛路演与答辩材料",
        "文档类型：系统理解与演示手册（论文式排版）",
    ]
    for line in meta:
        pm = doc.add_paragraph(line)
        set_paragraph_format(pm, first_line_indent=False)
        for run in pm.runs:
            apply_font_run(run, east_asia="宋体", size=12)

    pd = doc.add_paragraph(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(pd, first_line_indent=False)
    for run in pd.runs:
        apply_font_run(run, east_asia="宋体", size=12)

    doc.add_page_break()

    toc_h = doc.add_paragraph("目录")
    toc_h.style = "Heading 1"
    toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(toc_h, first_line_indent=False)

    toc_note = doc.add_paragraph("说明：为保证跨软件兼容，目录请在 WPS/Word 中右键“更新域”后自动生成。")
    set_paragraph_format(toc_note, first_line_indent=False)

    for line in lines:
        if line.startswith("## "):
            p_item = doc.add_paragraph(clean_inline(line[3:].strip()), style="List Number")
            set_paragraph_format(p_item, first_line_indent=False)

    doc.add_page_break()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        striped = line.strip()

        if not striped:
            doc.add_paragraph("")
            i += 1
            continue

        if striped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if striped.startswith("|") and "|" in striped[1:]:
            table_lines: list[str] = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith("|") and "|" in s[1:]:
                    table_lines.append(s)
                    i += 1
                else:
                    break

            if len(table_lines) >= 2:
                rows = [[clean_inline(x.strip()) for x in r.strip("|").split("|")] for r in table_lines]
                data_rows = [rows[0]]
                for r in rows[1:]:
                    if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in r):
                        continue
                    data_rows.append(r)

                col_count = max(len(r) for r in data_rows)
                t = doc.add_table(rows=len(data_rows), cols=col_count)
                t.style = "Table Grid"
                t.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row in enumerate(data_rows):
                    for c_idx in range(col_count):
                        value = row[c_idx] if c_idx < len(row) else ""
                        cell = t.cell(r_idx, c_idx)
                        cell.text = value
                        for cp in cell.paragraphs:
                            set_paragraph_format(cp, first_line_indent=False)
                            for rr in cp.runs:
                                apply_font_run(rr, east_asia="宋体", size=10.5)

                for cell in t.rows[0].cells:
                    for cp in cell.paragraphs:
                        for rr in cp.runs:
                            rr.bold = True

                doc.add_paragraph("")
            continue

        if line.startswith("# "):
            p_h1 = doc.add_paragraph(clean_inline(line[2:].strip()))
            p_h1.style = "Heading 1"
            p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p_h1, first_line_indent=False)
            i += 1
            continue

        if line.startswith("## "):
            p_h2 = doc.add_paragraph(clean_inline(line[3:].strip()))
            p_h2.style = "Heading 1"
            p_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p_h2, first_line_indent=False)
            i += 1
            continue

        if line.startswith("### "):
            p_h3 = doc.add_paragraph(clean_inline(line[4:].strip()))
            p_h3.style = "Heading 2"
            set_paragraph_format(p_h3, first_line_indent=False)
            i += 1
            continue

        if re.match(r"^\d+\.\s", striped):
            p_num = doc.add_paragraph(clean_inline(striped), style="List Number")
            set_paragraph_format(p_num, first_line_indent=False)
            i += 1
            continue

        if striped.startswith("- "):
            p_b = doc.add_paragraph(clean_inline(striped[2:]), style="List Bullet")
            set_paragraph_format(p_b, first_line_indent=False)
            i += 1
            continue

        p_text = doc.add_paragraph(clean_inline(line))
        set_paragraph_format(p_text, first_line_indent=True)
        i += 1

    for sec in doc.sections:
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)

    return doc


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    doc = build_document(text)
    doc.save(OUT_ASCII)
    OUT_CN.write_bytes(OUT_ASCII.read_bytes())
    print(str(OUT_ASCII))
    print(str(OUT_CN))


if __name__ == "__main__":
    main()

